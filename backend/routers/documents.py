import os
import shutil
import logging
import io
from typing import List
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, status, BackgroundTasks
from sqlalchemy.orm import Session

from backend.database import get_db
from backend.models import Document, DocumentChunk, KnowledgeNode, KnowledgeEdge
from backend.schemas import DocumentOut, DocumentStats
from backend.auth import get_current_user
from backend.ai.vector_store import vector_store_manager, chunk_text
from backend.ai.gemini_client import gemini_client

# Parsers
import pypdf
try:
    import docx
except ImportError:
    docx = None

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/documents", tags=["Documents"])

UPLOAD_DIR = "./uploaded_files"
os.makedirs(UPLOAD_DIR, exist_ok=True)

def parse_pdf(file_bytes: bytes) -> str:
    """Extracts text from PDF bytes."""
    pdf_file = io.BytesIO(file_bytes)
    reader = pypdf.PdfReader(pdf_file)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def parse_docx(file_bytes: bytes) -> str:
    """Extracts text from DOCX bytes."""
    if docx is None:
        logger.warning("python-docx is not installed. Unable to parse DOCX. Returning empty text.")
        return ""
    docx_file = io.BytesIO(file_bytes)
    doc = docx.Document(docx_file)
    text = []
    for para in doc.paragraphs:
        text.append(para.text)
    return "\n".join(text)

def process_document_pipeline(db: Session, doc_id: int, file_bytes: bytes):
    """Background task to parse, chunk, embed, index, and extract graph entities."""
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        return
        
    try:
        # 1. Parse File Content based on Type
        text = ""
        if doc.file_type == "PDF":
            text = parse_pdf(file_bytes)
        elif doc.file_type == "DOCX":
            text = parse_docx(file_bytes)
        else: # TXT
            text = file_bytes.decode("utf-8", errors="ignore")
            
        if not text.strip():
            raise ValueError("No text extracted from document")

        # 2. Chunk text
        chunks = chunk_text(text)
        if not chunks:
            raise ValueError("Document yielded 0 text chunks")
            
        # 3. Add to FAISS and get vector indices
        faiss_ids = vector_store_manager.add_chunks(chunks)
        
        # 4. Save chunks and indices in SQLite
        for idx, (chunk_text_content, faiss_id) in enumerate(zip(chunks, faiss_ids)):
            db_chunk = DocumentChunk(
                document_id=doc.id,
                chunk_index=idx,
                content=chunk_text_content,
                faiss_index_id=faiss_id
            )
            db.add(db_chunk)
            
        # 5. Extract Summary and Graph Entities using Gemini
        summary = ""
        if not gemini_client.is_demo_mode():
            # Get short summary
            try:
                summary_prompt = f"Summarize this industrial document in 2-3 sentences:\n\n{text[:4000]}"
                response = gemini_client.client.models.generate_content(
                    model="gemini-2.5-flash",
                    contents=summary_prompt
                )
                summary = response.text.strip()
            except Exception as e:
                logger.error(f"Failed to generate summary: {e}")
                summary = text[:200] + "..."
            
            # Auto extract graph elements
            try:
                graph_data = gemini_client.extract_graph_elements(text[:8000])
                for node in graph_data.get("nodes", []):
                    # Check if node already exists
                    existing_node = db.query(KnowledgeNode).filter(KnowledgeNode.id == node["id"]).first()
                    if not existing_node:
                        db_node = KnowledgeNode(
                            id=node["id"],
                            name=node["name"],
                            type=node["type"],
                            description=node.get("description", ""),
                            metadata_json=None
                        )
                        db.add(db_node)
                        
                db.commit() # Commit nodes before adding edges to avoid constraint issues
                
                for edge in graph_data.get("edges", []):
                    # Verify source and target exist in db
                    src = db.query(KnowledgeNode).filter(KnowledgeNode.id == edge["source"]).first()
                    tgt = db.query(KnowledgeNode).filter(KnowledgeNode.id == edge["target"]).first()
                    if src and tgt:
                        db_edge = KnowledgeEdge(
                            source_id=edge["source"],
                            target_id=edge["target"],
                            type=edge["type"],
                            description=edge.get("description", "")
                        )
                        db.add(db_edge)
            except Exception as e:
                logger.error(f"Failed to auto-extract graph elements: {e}")
        else:
            # Fallback mock summary
            summary = f"Summary of {doc.filename}. Contains industrial documentation relating to system operations."
            # Fallback mock graph elements
            graph_data = gemini_client._mock_extracted_graph(text)
            for node in graph_data.get("nodes", []):
                # Update document node ID
                if node["id"] == "DOC-CURR":
                    node["id"] = f"DOC-{doc.id}"
                existing_node = db.query(KnowledgeNode).filter(KnowledgeNode.id == node["id"]).first()
                if not existing_node:
                    db_node = KnowledgeNode(
                        id=node["id"],
                        name=node["name"],
                        type=node["type"],
                        description=node.get("description", ""),
                        metadata_json=None
                    )
                    db.add(db_node)
                    
            db.commit()
            
            for edge in graph_data.get("edges", []):
                if edge["source"] == "DOC-CURR":
                    edge["source"] = f"DOC-{doc.id}"
                if edge["target"] == "DOC-CURR":
                    edge["target"] = f"DOC-{doc.id}"
                src = db.query(KnowledgeNode).filter(KnowledgeNode.id == edge["source"]).first()
                tgt = db.query(KnowledgeNode).filter(KnowledgeNode.id == edge["target"]).first()
                if src and tgt:
                    db_edge = KnowledgeEdge(
                        source_id=edge["source"],
                        target_id=edge["target"],
                        type=edge["type"],
                        description=edge.get("description", "")
                    )
                    db.add(db_edge)
                    
        # Connect the document itself to the graph
        doc_node_id = f"DOC-{doc.id}"
        existing_doc_node = db.query(KnowledgeNode).filter(KnowledgeNode.id == doc_node_id).first()
        if not existing_doc_node:
            db_node = KnowledgeNode(
                id=doc_node_id,
                name=doc.filename,
                type="Document",
                description=f"Uploaded {doc.file_type} Document",
                metadata_json=None
            )
            db.add(db_node)
        
        doc.status = "Completed"
        doc.chunk_count = len(chunks)
        doc.summary = summary
        db.commit()
        logger.info(f"Pipeline completed for document ID {doc_id}: {doc.filename}")
        
    except Exception as e:
        logger.error(f"Pipeline failed for document ID {doc_id}: {e}")
        doc.status = "Failed"
        db.commit()

@router.post("/upload", response_model=DocumentOut, status_code=status.HTTP_201_CREATED)
async def upload_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    # Determine type
    filename = file.filename
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        file_type = "PDF"
    elif ext in [".docx", ".doc"]:
        file_type = "DOCX"
    elif ext in [".txt", ".md"]:
        file_type = "TXT"
    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Unsupported file format. Please upload PDF, DOCX, or TXT."
        )
        
    # Write file to disk temporarily
    file_path = os.path.join(UPLOAD_DIR, filename)
    # Check duplicate
    counter = 1
    base, extension = os.path.splitext(filename)
    while os.path.exists(file_path):
        file_path = os.path.join(UPLOAD_DIR, f"{base}_{counter}{extension}")
        filename = f"{base}_{counter}{extension}"
        counter += 1
        
    file_bytes = await file.read()
    with open(file_path, "wb") as f:
        f.write(file_bytes)
        
    # Insert metadata row into SQLite
    db_doc = Document(
        filename=filename,
        file_type=file_type,
        file_path=file_path,
        status="Processing",
        chunk_count=0
    )
    db.add(db_doc)
    db.commit()
    db.refresh(db_doc)
    
    # Process in background task
    background_tasks.add_task(process_document_pipeline, db, db_doc.id, file_bytes)
    
    return db_doc

@router.get("", response_model=List[DocumentOut])
def get_documents(db: Session = Depends(get_db), current_user = Depends(get_current_user)):
    return db.query(Document).order_by(Document.upload_date.desc()).all()

@router.delete("/{doc_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_document(doc_id: int, db: Session = Depends(get_db), current_user = Depends(get_current_user)):
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Document not found"
        )
        
    # Delete from filesystem
    if os.path.exists(doc.file_path):
        try:
            os.remove(doc.file_path)
        except Exception as e:
            logger.error(f"Failed to remove file from disk: {e}")
            
    # Delete Document node from Knowledge Graph
    doc_node_id = f"DOC-{doc.id}"
    db.query(KnowledgeNode).filter(KnowledgeNode.id == doc_node_id).delete()
    
    # Delete from DB (cascade handles chunks & findings)
    db.delete(doc)
    db.commit()
    return

@router.get("/stats/telemetry", response_model=DocumentStats)
def get_document_stats(db: Session = Depends(get_db), current_user = Depends(get_current_user)):
    total_docs = db.query(Document).count()
    total_chunks = db.query(DocumentChunk).count()
    processing_count = db.query(Document).filter(Document.status == "Processing").count()
    failed_count = db.query(Document).filter(Document.status == "Failed").count()
    
    return {
        "total_documents": total_docs,
        "total_chunks": total_chunks,
        "processing_count": processing_count,
        "failed_count": failed_count
    }
